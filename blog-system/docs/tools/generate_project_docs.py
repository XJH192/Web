from __future__ import annotations

import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[3]
DOCS_DIR = ROOT / "blog-system" / "docs"
OUTPUT_DIR = DOCS_DIR / "docx"
PROJECT_NAME = "Ciallo～(∠・ω< )⌒☆ 多用户智能博客系统"
PROJECT_TEAM = "Ciallo～(∠・ω< )⌒☆ 项目组"
PROJECT_FILENAME_PREFIX = "Ciallo多用户智能博客系统"

INK = "26364A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
MUTED = "667085"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F4F6F9"
BORDER = "C9D2DE"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = 120


def set_run_font(run, size=None, bold=None, italic=None, color=None, mono=False):
    ascii_font = "Consolas" if mono else "Calibri"
    east_asia = "Microsoft YaHei"
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, color=BORDER, size="8", space="5"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    for edge in ("top", "left", "bottom", "right"):
        node = p_bdr.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            p_bdr.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), space)
        node.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=CELL_MARGIN_DXA, bottom=80, end=CELL_MARGIN_DXA):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_repeat_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            set_cell_margins(cell)


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_end)
    suffix = paragraph.add_run(" 页")
    set_run_font(suffix, size=9, color=MUTED)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    heading_specs = {
        "Heading 1": (16, BLUE, 18, 10),
        "Heading 2": (13, BLUE, 14, 7),
        "Heading 3": (12, DARK_BLUE, 10, 5),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code_style = doc.styles.add_style("Code Block", 1)
    code_style.font.name = "Consolas"
    code_style._element.rPr.rFonts.set(qn("w:ascii"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:hAnsi"), "Consolas")
    code_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    code_style.font.size = Pt(9)
    code_style.font.color.rgb = RGBColor.from_string(INK)
    code_style.paragraph_format.left_indent = Inches(0.12)
    code_style.paragraph_format.right_indent = Inches(0.12)
    code_style.paragraph_format.space_before = Pt(5)
    code_style.paragraph_format.space_after = Pt(8)
    code_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    callout = doc.styles.add_style("Callout", 1)
    callout.font.name = "Calibri"
    callout._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    callout.font.size = Pt(10.5)
    callout.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    callout.paragraph_format.left_indent = Inches(0.14)
    callout.paragraph_format.right_indent = Inches(0.14)
    callout.paragraph_format.space_before = Pt(5)
    callout.paragraph_format.space_after = Pt(8)
    callout.paragraph_format.line_spacing = 1.2


def add_numbering(doc, bullet):
    numbering = doc.part.numbering_part.element
    existing_ids = [int(x.get(qn("w:abstractNumId"))) for x in numbering.findall(qn("w:abstractNum"))]
    abstract_id = max(existing_ids or [0]) + 1
    existing_num_ids = [int(x.get(qn("w:numId"))) for x in numbering.findall(qn("w:num"))]
    num_id = max(existing_num_ids or [0]) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "multilevel")
    abstract.append(multi)

    for level in range(3):
        lvl = OxmlElement("w:lvl")
        lvl.set(qn("w:ilvl"), str(level))
        start = OxmlElement("w:start")
        start.set(qn("w:val"), "1")
        lvl.append(start)
        num_fmt = OxmlElement("w:numFmt")
        num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
        lvl.append(num_fmt)
        lvl_text = OxmlElement("w:lvlText")
        if bullet:
            lvl_text.set(qn("w:val"), ["●", "○", "■"][level])
        else:
            lvl_text.set(qn("w:val"), f"%{level + 1}.")
        lvl.append(lvl_text)
        lvl_jc = OxmlElement("w:lvlJc")
        lvl_jc.set(qn("w:val"), "left")
        lvl.append(lvl_jc)
        p_pr = OxmlElement("w:pPr")
        tabs = OxmlElement("w:tabs")
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "num")
        tab.set(qn("w:pos"), str(540 + level * 360))
        tabs.append(tab)
        p_pr.append(tabs)
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(540 + level * 360))
        ind.set(qn("w:hanging"), "270")
        p_pr.append(ind)
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:after"), "80")
        spacing.set(qn("w:line"), "300")
        spacing.set(qn("w:lineRule"), "auto")
        p_pr.append(spacing)
        lvl.append(p_pr)
        abstract.append(lvl)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_num_id = OxmlElement("w:abstractNumId")
    abstract_num_id.set(qn("w:val"), str(abstract_id))
    num.append(abstract_num_id)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id, level):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        p_pr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(max(0, min(level, 2))))
    num_id_node = OxmlElement("w:numId")
    num_id_node.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl)
    num_pr.append(num_id_node)


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^)]+\))")


def add_inline(paragraph, text, size=None, color=None):
    pos = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, color=color)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, bold=True, color=color)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, size=size or 10, color=DARK_BLUE, mono=True)
            set_paragraph_shading(paragraph, "F7F8FA")
        else:
            label = token[1:token.index("](")]
            url = token[token.index("](") + 2:-1]
            run = paragraph.add_run(f"{label}（{url}）")
            set_run_font(run, size=size, color=BLUE)
            run.underline = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, color=color)


def parse_table(lines, index):
    rows = []
    while index < len(lines) and lines[index].strip().startswith("|"):
        raw = lines[index].strip().strip("|")
        rows.append([cell.strip() for cell in raw.split("|")])
        index += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    return rows, index


def column_widths(rows):
    cols = max(len(row) for row in rows)
    weights = []
    for col in range(cols):
        values = [row[col] if col < len(row) else "" for row in rows]
        max_len = max([len(re.sub(r"[`*]", "", value)) for value in values] + [6])
        weights.append(max(7, min(max_len, 46)))
    if cols >= 3:
        for idx, row in enumerate(rows[:2]):
            if row and row[0] in ("编号", "软件", "成员", "表名"):
                weights[0] = min(weights[0], 10)
    total = sum(weights)
    widths = [max(900, int(TABLE_WIDTH_DXA * value / total)) for value in weights]
    while sum(widths) > TABLE_WIDTH_DXA:
        widest = max(range(len(widths)), key=lambda i: widths[i])
        if widths[widest] <= 900:
            break
        widths[widest] -= 1
    while sum(widths) < TABLE_WIDTH_DXA:
        widest = max(range(len(widths)), key=lambda i: weights[i])
        widths[widest] += 1
    return widths


def add_table(doc, rows):
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = column_widths(rows)
    set_repeat_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    for row_idx, values in enumerate(rows):
        row = table.rows[row_idx]
        for col_idx in range(cols):
            cell = row.cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.15
            value = values[col_idx] if col_idx < len(values) else ""
            add_inline(paragraph, value, size=9.2)
            if row_idx == 0:
                set_cell_shading(cell, LIGHT_BLUE)
                for run in paragraph.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif len(value) <= 12 and col_idx < 2:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(2)


def add_code_block(doc, lines):
    clean = "\n".join(lines).rstrip()
    if clean.startswith("[在此插入截图"):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(30)
        paragraph.paragraph_format.space_after = Pt(30)
        set_paragraph_shading(paragraph, "F8FAFC")
        set_paragraph_border(paragraph, BORDER, "10", "8")
        run = paragraph.add_run(clean)
        set_run_font(run, size=10, bold=True, color=MUTED)
        return
    paragraph = doc.add_paragraph(style="Code Block")
    set_paragraph_shading(paragraph, LIGHT_GRAY)
    set_paragraph_border(paragraph, BORDER, "4", "4")
    for idx, line in enumerate(lines):
        run = paragraph.add_run(line)
        set_run_font(run, size=9, mono=True, color=INK)
        if idx < len(lines) - 1:
            run.add_break()


def add_cover(doc, title, subtitle, doc_type):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header.paragraph_format.space_after = Pt(0)
    run = header.add_run(f"{PROJECT_NAME}  |  {doc_type}")
    set_run_font(run, size=9, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    add_page_field(footer)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(105)
    spacer.paragraph_format.space_after = Pt(0)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(16)
    run = kicker.add_run("课程设计交付文档")
    set_run_font(run, size=11, bold=True, color=BLUE)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(title)
    set_run_font(run, size=28, bold=True, color=INK)

    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run(subtitle)
    set_run_font(run, size=14, color=DARK_BLUE)

    rule = doc.add_paragraph()
    rule.paragraph_format.left_indent = Inches(1.3)
    rule.paragraph_format.right_indent = Inches(1.3)
    rule.paragraph_format.space_after = Pt(40)
    p_pr = rule._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BLUE)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

    for label, value in (
        ("项目名称", PROJECT_NAME),
        ("技术栈", "Hexo 7 + Spring Boot 2.7 + MySQL 8 + DeepSeek"),
        ("小组规模", "3 人"),
        ("文档版本", "V1.0"),
        ("编制日期", date.today().strftime("%Y 年 %m 月 %d 日")),
    ):
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(4)
        label_run = paragraph.add_run(label + "：")
        set_run_font(label_run, size=10.5, bold=True, color=MUTED)
        value_run = paragraph.add_run(value)
        set_run_font(value_run, size=10.5, color=INK)

    doc.add_page_break()


def add_contents(doc, headings):
    paragraph = doc.add_paragraph("目录", style="Heading 1")
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.page_break_before = False
    intro = doc.add_paragraph("以下目录对应 Word 导航窗格中的一级标题，可直接使用标题导航浏览。")
    intro.paragraph_format.space_after = Pt(10)
    num_id = add_numbering(doc, bullet=False)
    for heading in headings:
        paragraph = doc.add_paragraph()
        apply_numbering(paragraph, num_id, 0)
        add_inline(paragraph, heading)


def markdown_to_docx(source_path, output_path, doc_type):
    text = source_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    title = re.sub(r"^#\s+", "", lines[0]).strip()
    headings = [re.sub(r"^##\s+", "", line).strip() for line in lines if line.startswith("## ")]

    doc = Document()
    style_document(doc)
    doc.core_properties.title = title
    doc.core_properties.subject = f"{PROJECT_NAME}{doc_type}"
    doc.core_properties.author = PROJECT_TEAM
    doc.core_properties.keywords = "Hexo, Spring Boot, MySQL, DeepSeek, 博客系统"
    doc.core_properties.comments = "由项目 Markdown 文档生成"

    add_cover(
        doc,
        title.replace(PROJECT_NAME, "").strip() or doc_type,
        PROJECT_NAME,
        doc_type,
    )
    add_contents(doc, headings)

    bullet_num_id = add_numbering(doc, bullet=True)
    decimal_num_id = add_numbering(doc, bullet=False)
    index = 1
    paragraph_buffer = []
    in_code = False
    code_lines = []

    def flush_paragraph():
        nonlocal paragraph_buffer
        if paragraph_buffer:
            paragraph = doc.add_paragraph()
            add_inline(paragraph, " ".join(part.strip() for part in paragraph_buffer))
            paragraph_buffer = []

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                code_lines.append(line)
            index += 1
            continue

        if stripped.startswith("```"):
            flush_paragraph()
            in_code = True
            code_lines = []
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            index += 1
            continue

        if stripped.startswith("|"):
            flush_paragraph()
            rows, index = parse_table(lines, index)
            add_table(doc, rows)
            continue

        heading_match = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading_match:
            flush_paragraph()
            source_level = len(heading_match.group(1))
            word_level = min(3, source_level - 1)
            paragraph = doc.add_paragraph(heading_match.group(2), style=f"Heading {word_level}")
            if word_level == 1:
                paragraph.paragraph_format.page_break_before = True
            index += 1
            continue

        if stripped.startswith(">"):
            flush_paragraph()
            paragraph = doc.add_paragraph(style="Callout")
            set_paragraph_shading(paragraph, LIGHT_BLUE)
            set_paragraph_border(paragraph, "AFC5DD", "6", "5")
            add_inline(paragraph, stripped.lstrip(">").strip(), size=10.5, color=DARK_BLUE)
            index += 1
            continue

        unordered = re.match(r"^(\s*)[-*]\s+(.+)$", line)
        ordered = re.match(r"^(\s*)\d+[.)]\s+(.+)$", line)
        if unordered or ordered:
            flush_paragraph()
            match = unordered or ordered
            level = min(2, len(match.group(1).replace("\t", "    ")) // 2)
            if ordered:
                previous = lines[index - 1] if index > 0 else ""
                if not re.match(r"^\s*\d+[.)]\s+(.+)$", previous):
                    decimal_num_id = add_numbering(doc, bullet=False)
            paragraph = doc.add_paragraph()
            apply_numbering(paragraph, bullet_num_id if unordered else decimal_num_id, level)
            add_inline(paragraph, match.group(2))
            index += 1
            continue

        if re.fullmatch(r"-{3,}", stripped):
            flush_paragraph()
            paragraph = doc.add_paragraph()
            p_pr = paragraph._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "4")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), BORDER)
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
            index += 1
            continue

        paragraph_buffer.append(stripped)
        index += 1

    flush_paragraph()
    if in_code:
        add_code_block(doc, code_lines)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def main():
    outputs = [
        (
            DOCS_DIR / "需求分析文档.md",
            OUTPUT_DIR / f"{PROJECT_FILENAME_PREFIX}-需求分析文档.docx",
            "需求分析文档",
        ),
        (
            DOCS_DIR / "系统使用手册.md",
            OUTPUT_DIR / f"{PROJECT_FILENAME_PREFIX}-系统使用手册.docx",
            "系统使用手册",
        ),
    ]
    for source, target, doc_type in outputs:
        markdown_to_docx(source, target, doc_type)
        print(target)


if __name__ == "__main__":
    sys.exit(main())
